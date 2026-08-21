# -*- coding: utf-8 -*-
"""Genera un documento Word (.docx) con el resumen de un caso de consulta
horaria: datos de la carta, significadores, y las preguntas/respuestas."""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def generar_docx(informe, historial, pregunta_original="", meta_extra=None):
    doc = Document()

    title = doc.add_heading("Consulta de Astrología Horaria — Método Ben Ezra", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cm = informe.get("chart_meta", {})
    p = doc.add_paragraph()
    p.add_run(f"Fecha/hora: {cm.get('fecha_local', '—').replace('T', ' ')}\n").bold = False
    p.add_run(f"Ascendente: {cm.get('asc', '—')}    Medio Cielo: {cm.get('mc', '—')}\n")
    p.add_run(f"Carta: {'Diurna' if cm.get('carta_diurna') else 'Nocturna'}    ")
    p.add_run(f"Regente de la Hora: {cm.get('regente_hora', '—')}\n")
    if pregunta_original:
        p.add_run(f"\nPregunta original: ").bold = True
        p.add_run(pregunta_original)

    _heading(doc, "Significadores", level=1)
    sig = informe.get("significadores", {})
    if sig.get("consultante"):
        c = sig["consultante"]
        doc.add_paragraph(
            f"Consultante (I): {c.get('significador_recomendado', '—')} — "
            f"regente cúspide: {c.get('regente_cuspide', '—')}, "
            f"almuten(es): {', '.join(c.get('almutenes_cuspide', []))}"
        )
        doc.add_paragraph(c.get("razonamiento", ""), style=None)
    if sig.get("pregunta"):
        q = sig["pregunta"]
        doc.add_paragraph(
            f"Pregunta (casa {q.get('casa')}): {q.get('significador_recomendado', '—')} — "
            f"regente cúspide: {q.get('regente_cuspide', '—')}, "
            f"almuten(es): {', '.join(q.get('almutenes_cuspide', []))}"
        )
        doc.add_paragraph(q.get("razonamiento", ""))

    _heading(doc, "Planetas", level=1)
    planetas = informe.get("planetas", {})
    if planetas:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Planeta", "Signo", "Casa", "Dignidad"
        for nombre, d in planetas.items():
            row = table.add_row().cells
            row[0].text = nombre + (" ℞" if d.get("retrógrado") else "")
            row[1].text = f"{d.get('signo','')} {d.get('grado','')}°"
            row[2].text = str(d.get("casa", ""))
            row[3].text = ", ".join(d.get("dignidad_esencial", []))

    luna = informe.get("luna", {})
    _heading(doc, "Luna", level=1)
    doc.add_paragraph(f"Vacía de curso: {'Sí' if luna.get('vacia_de_curso') else 'No'}")
    doc.add_paragraph(f"Vía combusta: {'Sí' if luna.get('via_combusta') else 'No'}")
    if luna.get("ultimo_aspecto"):
        u = luna["ultimo_aspecto"]
        doc.add_paragraph(f"Último aspecto: {u.get('aspecto')} con {u.get('planeta')} (hace {abs(u.get('dias',0))} días)")
    if luna.get("proximo_aspecto"):
        n = luna["proximo_aspecto"]
        doc.add_paragraph(f"Próximo aspecto: {n.get('aspecto')} con {n.get('planeta')} (en {n.get('dias')} días)")

    validez = informe.get("validez_tema", {})
    _heading(doc, "Validez del tema", level=1)
    ah = validez.get("almuten_hora", {})
    doc.add_paragraph(f"Almuten I / Regente Hora: {'Coinciden' if ah.get('coincide') else 'Sin coincidencia'}")
    for nota in ah.get("notas", []):
        doc.add_paragraph(f"• {nota}", style=None)
    if validez.get("saturno_en_vii"):
        doc.add_paragraph("• Saturno está en la casa VII.")

    if historial:
        _heading(doc, "Preguntas y respuestas", level=1)
        for i, qa in enumerate(historial, 1):
            preg = doc.add_paragraph()
            preg.add_run(f"P{i}: {qa.get('pregunta', '')}").bold = True
            doc.add_paragraph(qa.get("respuesta", ""))
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
